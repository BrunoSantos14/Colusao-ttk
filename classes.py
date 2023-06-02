import pandas as pd
from docx import Document
import datetime as dt
import os
import win32com.client
from subprocess import Popen
import tkinter as tk


def criar_col_item(df):
    """Retorna uma tupla de 3 elementos, com o dataframe adicionado de uma coluna informando o item
    (ano + rodada + num_item), além da rodada inicial e rodada final."""
    df = df.copy(deep=True)  # Para evitar warnings no código
    # Mudando formatação para data e criando colunas Ano e Mes
    dic_meses = {'Jan': '1', 'Fev': '2', 'Mar': '3', 'Abr': '4', 'Mai': '5', 'Jun': '6', 'Jul': '7', 'Ago': '8',
                    'Set': '9', 'Out': '10', 'Nov': '11', 'Dez': '12'}

    df['NOME_ENVIO'] = df['NOME_ENVIO'].apply(lambda x: x.replace(x[0:3], dic_meses[x[0:3]]))
    df['NOME_ENVIO'] = pd.to_datetime(df['NOME_ENVIO'], format='%m/%Y')

    rodada_min = min(df['NOME_ENVIO'])  # Selecionando menor data (será colocado no df final)
    rodada_max = max(df['NOME_ENVIO'])  # Selecionando maior data (será colocado no df final)

    df['Ano'] = df['NOME_ENVIO'].dt.year  # Criando uma coluna de ano
    df['Mes'] = df['NOME_ENVIO'].dt.month  # Criando uma coluna de mês

    # Criando a coluna Rodada
    df.loc[df['Mes'].isin([10, 11, 12]), 'RODADA'] = '1'
    df.loc[df['Mes'].isin([1, 2, 3]), 'RODADA'] = '2'
    df.loc[df['Mes'].isin([4, 5, 6]), 'RODADA'] = '3'
    df.loc[df['Mes'].isin([7, 8, 9]), 'RODADA'] = '4'

    # Deixando a coluna Ano apenas com dois dígitos
    df.loc[df['RODADA'] == '1', 'Ano'] += 1  # Colocando itens da primeira rodada para o próximo ano
    df['Ano'] = df['Ano'].astype(str).apply(lambda x: x[-2:])

    # Criando a coluna ITEM (as linhas vão virar colunas)
    df['NUM_ITEM'] = df.NUM_ITEM.astype(str)
    df['ITEM'] = 'A' + df['Ano'] + 'R' + df['RODADA'] + 'I' + df['NUM_ITEM']

    df = df.drop(columns=['Ano', 'Mes', 'RODADA'], axis=1)
    return df, rodada_min, rodada_max


class Colas:
    def __init__(self, ano):
        self.ano = ano


    def listar_colas(self, df):
        """Quebra os resultados informados em colunas por item de rodada e filtra apenas resultados idênticos, gerando
        suspeitos de colusão. Estes participantes são listados conjuntamente, com a informação da quantidade de
        determinações observadas."""

        df, rodada_min, rodada_max = criar_col_item(df)

        df = df[['PART', 'MODULO', 'NOME_DET', 'NOME_ENVIO', 'NUM_ITEM', 'VALOR', 'ITEM']]
        df.columns = ['Cliente', 'MODULO', 'Analito', 'NOME_ENVIO', 'NUM_ITEM', 'VALOR', 'ITEM']

        # Mantendo um valor por participante em cada rodada
        df = df.copy(deep=True)
        df['VALOR'] = pd.to_numeric(df['VALOR'])
        df['VALOR'] = df.groupby(['Cliente', 'MODULO', 'Analito', 'ITEM'])['VALOR'].transform('median')
        df.drop_duplicates(subset=['Cliente', 'MODULO', 'Analito', 'ITEM'],
                           inplace=True,
                           keep='first')

        # Escolhendo as colunas necessárias
        df = df[['Cliente', 'MODULO', 'Analito', 'VALOR', 'ITEM']]

        # Transformando os valores de ITEM para colunas
        df["constante"] = 10
        df = df.pivot_table(
            values="VALOR",
            index=['Cliente', 'MODULO', 'Analito'],
            columns='ITEM',
        )
        df = df.reset_index()

        # Filtrando apenas 60% de resultados reportados por participante
        qtd_colunas = len(df.columns) - 3
        df['QTD_preenchido'] = df.count(axis=1) - 3
        df['Perc'] = df['QTD_preenchido'] / qtd_colunas
        df = df.loc[df['Perc'] >= 0.6, :]

        df.drop(['QTD_preenchido', 'Perc'], axis=1, inplace=True)

        if len(df) == 0:
            # return pd.DataFrame({'Conclusão': ['Ninguém com mais de 60% preenchido']})
            return 'Ninguém com mais de 60% preenchido'

        # Contando quantos dados de cada analito existem. Filtrando para 3 dados ou mais por analito
        for analito in list(df['Analito'].unique()):
            df.loc[df['Analito'] == analito, 'QTD_dados'] = len(df.loc[df['Analito'] == analito])

        df = df.loc[df['QTD_dados'] >= 3, :]
        df = df.drop('QTD_dados', axis=1)

        # Mantendo apenas linhas onde os dados são repetidos
        bool_series = df.duplicated(subset=list(df.columns[1:]), keep=False)
        df = df[bool_series]

        # Apenas para não retornar df vazio e não gerar erro
        if len(df) == 0:
            # return pd.DataFrame({'Conclusão': ['Sem resultados idênticos no ano']})
            return 'Sem resultados idênticos no ano'

        # Colocando as colas uma embaixo da outra
        df.sort_values(by=list(df.columns[1:]) + ['Cliente'], ascending=True, inplace=True)

        # Calculando quantidade de analitos, quantidade de colas por grupos de participantes e calculando percentual
        df['Qtd_cola'] = df.groupby('Cliente')['Cliente'].transform('count')

        # Criando a coluna grupo para juntar clientes em colusão
        df['Cliente'] = df.Cliente.astype(str)
        df.fillna('', inplace=True)

        colunas_itens = [i for i in df.columns if str(self.ano)[:-2] in i]
        for coluna in colunas_itens:
            df[coluna] = pd.to_numeric(df[coluna])
            # df[coluna].fillna(np.nan, inplace=True)

        aux = df.groupby(by=list(df.columns[1:-2]))['Cliente'].agg(lambda col: ' - '.join(col)).reset_index()

        df = df.merge(aux, on=list(df.columns[1:-2]), suffixes=("", "_x"))
        df.rename(columns={'Cliente_x': 'Grupos'}, inplace=True)

        df['Cliente'] = pd.to_numeric(df['Cliente'])#.astype(str).str.replace(',', '')
        return df, rodada_min, rodada_max

    def lista_cola_resum(self, df):
        df = df[['Grupos', 'Qtd_cola']]
        df = df.drop_duplicates()
        return df.sort_values(by='Qtd_cola', ascending=False).reset_index(drop=True)

    def filtrar_analitos(self, df):
        df = df[['Analito']]
        df = df.drop_duplicates()
        return df.sort_values(by='Analito', ascending=False).reset_index(drop=True)


class ModeloColusao:

    def __init__(self, path: str, resp: str, id_modulo, modulo, analise: str, 
               colusao: str, investigados, selecionados, cadastrado, retirados):
        """
        path: Caminho do arquivo a ser lido
        resp: Responsável pelo relatório
        id_modulo: id do modulo do estudo
        analise: Exames Analisados
        colusao: Exames em Colusão
        investigados: Participantes Investigados
        selecionados: Participantes Selecionados
        cadastro: Grupos cadastrados
        retirados: Grupos retirados de cadastro
        """    
        self.__doc = Document(path)
        self.__resp = resp
        self.__doc = self.__cabecalho()
        self.__id = str(id_modulo)
        self.__mod = str(modulo)
        self.__analise = str(analise)
        self.__colusao = str(colusao)
        self.__inv = str(investigados)
        self.__sel = str(selecionados)
        self.__cad = str(cadastrado)
        self.__ret = str(retirados)
        self.__doc = self.__corpo()

    def __cabecalho(self):
        meses = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
                 'julho', 'agosto', 'setembro', 'outubro', 'novembro',
                 'dezembro']

        self.__doc.paragraphs[1].text = self.__doc.paragraphs[1]\
                                           .text.replace('<dia>',
                                                         str(dt.datetime.now().day)
                                                         )

        self.__doc.paragraphs[1].text = self.__doc.paragraphs[1]\
                                                    .text.replace('<mes>', meses[dt.datetime.now().month-1])

        self.__doc.paragraphs[1].text = self.__doc.paragraphs[1]\
                                                    .text.replace('<ano>', str(dt.datetime.now().year))

        self.__doc.paragraphs[3].text = self.__doc.paragraphs[3]\
                                                    .text.replace('<responsavel>', self.__resp)
        return self.__doc

    def __corpo(self):

        self.__doc.paragraphs[9].text = self.__doc.paragraphs[9]\
                                                    .text.replace('<id_modulo>', self.__id)

        self.__doc.paragraphs[10].text = self.__doc.paragraphs[10]\
                                                    .text.replace('<modulo>', self.__mod)

        self.__doc.paragraphs[13].text = self.__doc.paragraphs[13]\
                                                    .text.replace('<analisado>', self.__analise)

        self.__doc.paragraphs[14].text = self.__doc.paragraphs[14]\
                                                    .text.replace('<colusao>', self.__colusao)

        self.__doc.paragraphs[15].text = self.__doc.paragraphs[15]\
                                                    .text.replace('<investigados>', self.__inv)

        self.__doc.paragraphs[16].text = self.__doc.paragraphs[16]\
                                                    .text.replace('<selecionados>', self.__sel)

        self.__doc.paragraphs[21].text = self.__cad

        self.__doc.paragraphs[25].text = self.__ret
        return self.__doc

    def salvar(self, path, pasta: str, mode='pdf', app='word'):
        '''path: Caminho para salvar o novo arquivo
           mode: formato de saída do arquivo ['word', 'pdf']
           app: 'word' ou 'libreoffice'
           pasta: Pasta onde será armazenado o pdf
        '''

        self.__doc.save(pasta+f'\\{path}.docx')

        if mode.lower() == 'word':
            return "Salvo em word"

        elif (mode.lower() == 'pdf') & (app.lower() == 'word'):
            wdFormatPDF = 17

            word = win32com.client.Dispatch('Word.Application')
            doc = word.Documents.Open(pasta+f'\\{path}.docx')
            doc.SaveAs(pasta+f'\\{path}.pdf', FileFormat=wdFormatPDF)
            doc.Close()
            word.Quit()

            os.remove(pasta+f'\\{path}.docx')
            return "Salvo em pdf"

        elif (mode.lower() == 'pdf') & (app.lower() == 'libreoffice'):

            LIBRE_OFFICE = r'C:\Program Files\LibreOffice\program\soffice.exe'

            arquivo_de_entrada = pasta+f'\\{path}.docx'

            p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir', pasta, arquivo_de_entrada])
            p.communicate()
            os.remove(pasta+f'\\{path}.docx')
            return "Salvo em pdf"
        return "Error"
    


class CreateToolTip(object):
    """
    create a tooltip for a given widget
    """
    def __init__(self, widget, text='widget info'):
        self.waittime = 500     #miliseconds
        self.wraplength = 180   #pixels
        self.widget = widget
        self.text = text
        self.widget.bind("<Enter>", self.enter)
        self.widget.bind("<Leave>", self.leave)
        self.widget.bind("<ButtonPress>", self.leave)
        self.id = None
        self.tw = None

    def enter(self, event=None):
        self.schedule()

    def leave(self, event=None):
        self.unschedule()
        self.hidetip()

    def schedule(self):
        self.unschedule()
        self.id = self.widget.after(self.waittime, self.showtip)

    def unschedule(self):
        id = self.id
        self.id = None
        if id:
            self.widget.after_cancel(id)

    def showtip(self, event=None):
        x = y = 0
        x, y, cx, cy = self.widget.bbox("insert")
        x += self.widget.winfo_rootx() + 25
        y += self.widget.winfo_rooty() + 20
        # creates a toplevel window
        self.tw = tk.Toplevel(self.widget)
        # Leaves only the label and removes the app window
        self.tw.wm_overrideredirect(True)
        self.tw.wm_geometry("+%d+%d" % (x, y))
        label = tk.Label(self.tw, text=self.text, justify='left',
                       background="#ffffff", relief='solid', borderwidth=1,
                       wraplength = self.wraplength)
        label.pack(ipadx=1)

    def hidetip(self):
        tw = self.tw
        self.tw= None
        if tw:
            tw.destroy()